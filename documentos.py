from datetime import datetime
from pathlib import Path
import json

from docx import Document
from pypdf import PdfReader

from config import ARQUIVO_STATUS, FORMATOS_ACEITOS, PASTA_UPLOADS


def carregar_desativados():
    if not ARQUIVO_STATUS.exists():
        return set()

    try:
        return set(json.loads(ARQUIVO_STATUS.read_text(encoding="utf-8")))
    except Exception:
        return set()


def salvar_desativados(desativados):
    ARQUIVO_STATUS.write_text(
        json.dumps(sorted(desativados), ensure_ascii=False),
        encoding="utf-8",
    )


def listar_documentos():
    desativados = carregar_desativados()
    documentos = []

    for arquivo in PASTA_UPLOADS.iterdir():
        if arquivo.is_file() and arquivo.suffix.lower() in FORMATOS_ACEITOS:
            documentos.append(
                {
                    "nome": arquivo.name,
                    "tipo": arquivo.suffix.upper().replace(".", ""),
                    "data": datetime.fromtimestamp(
                        arquivo.stat().st_mtime
                    ).strftime("%d/%m/%Y %H:%M"),
                    "ativo": arquivo.name not in desativados,
                }
            )

    return sorted(
        documentos,
        key=lambda documento: documento["data"],
        reverse=True,
    )


def documentos_ativos():
    desativados = carregar_desativados()

    return sorted(
        [
            arquivo
            for arquivo in PASTA_UPLOADS.iterdir()
            if (
                arquivo.is_file()
                and arquivo.suffix.lower() in FORMATOS_ACEITOS
                and arquivo.name not in desativados
            )
        ],
        key=lambda arquivo: arquivo.stat().st_mtime,
    )


def extrair_texto(caminho):
    extensao = caminho.suffix.lower()

    if extensao == ".txt":
        return caminho.read_text(encoding="utf-8", errors="replace")

    if extensao == ".docx":
        documento = Document(str(caminho))
        partes = [paragrafo.text for paragrafo in documento.paragraphs]

        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.append(
                    " | ".join(celula.text for celula in linha.cells)
                )

        return "\n".join(partes)

    if extensao == ".pdf":
        leitor = PdfReader(str(caminho))
        return "\n".join(
            pagina.extract_text() or ""
            for pagina in leitor.pages
        )

    raise ValueError("Formato de arquivo não aceito.")


def extrair_partes_referenciadas(caminho):
    extensao = caminho.suffix.lower()

    if extensao == ".pdf":
        leitor = PdfReader(str(caminho))
        partes = []

        for numero_pagina, pagina in enumerate(leitor.pages, start=1):
            texto = pagina.extract_text() or ""

            if texto.strip():
                partes.append(
                    {
                        "pagina": numero_pagina,
                        "texto": texto,
                    }
                )

        return partes

    return [
        {
            "pagina": None,
            "texto": extrair_texto(caminho),
        }
    ]