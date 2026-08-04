from datetime import datetime
from pathlib import Path
import html
import json
import os
import shutil
import math

from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.responses import HTMLResponse
from openai import OpenAI
from pypdf import PdfReader

load_dotenv(".env.local")

app = FastAPI()

PASTA_UPLOADS = Path("uploads")
PASTA_UPLOADS.mkdir(exist_ok=True)

ARQUIVO_STATUS = Path("documentos_desativados.json")
FORMATOS_ACEITOS = {".txt", ".docx", ".pdf"}
ARQUIVO_INDICE = Path("indice_documentos.json")
TAMANHO_TRECHO = 1200
SOBREPOSICAO = 200
QUANTIDADE_TRECHOS = 6


def carregar_desativados():
    if not ARQUIVO_STATUS.exists():
        return set()

    try:
        return set(json.loads(ARQUIVO_STATUS.read_text(encoding="utf-8")))
    except Exception:
        return set()


def salvar_desativados(desativados):
    ARQUIVO_STATUS.write_text(
        json.dumps(sorted(desativados), ensure_ascii=False),
        encoding="utf-8",
    )


def listar_documentos():
    desativados = carregar_desativados()
    documentos = []

    for arquivo in PASTA_UPLOADS.iterdir():
        if arquivo.is_file() and arquivo.suffix.lower() in FORMATOS_ACEITOS:
            documentos.append(
                {
                    "nome": arquivo.name,
                    "tipo": arquivo.suffix.upper().replace(".", ""),
                    "data": datetime.fromtimestamp(
                        arquivo.stat().st_mtime
                    ).strftime("%d/%m/%Y %H:%M"),
                    "ativo": arquivo.name not in desativados,
                }
            )

    return sorted(documentos, key=lambda documento: documento["data"], reverse=True)


def documentos_html():
    documentos = listar_documentos()

    if not documentos:
        return """
        <div class="documentos">
            <h2>Documentos enviados</h2>
            <p>Nenhum documento enviado ainda.</p>
        </div>
        """

    itens = []

    for documento in documentos:
        nome = html.escape(documento["nome"])
        estado = "Ativo" if documento["ativo"] else "Desativado"
        acao = "Desativar" if documento["ativo"] else "Reativar"

        itens.append(
            f"""
            <div class="documento">
                <div>
                    <strong>{nome}</strong><br>
                    <span>{documento["tipo"]} · {documento["data"]} · {estado}</span>
                </div>
                <div>
                    <form action="/alternar" method="post" class="acao">
                        <input type="hidden" name="nome" value="{nome}">
                        <button type="submit" class="secundario">{acao}</button>
                    </form>
                    <form action="/remover" method="post" class="acao"
                          onsubmit="return confirm('Remover este documento?')">
                        <input type="hidden" name="nome" value="{nome}">
                        <button type="submit" class="remover">Remover</button>
                    </form>
                </div>
            </div>
            """
        )

    return f"""
    <div class="documentos">
        <h2>Documentos enviados</h2>
        {''.join(itens)}
                <form action="/reindexar" method="post">
            <button type="submit">Preparar biblioteca para consulta</button>
        </form>
    </div>
    """


def pagina(mensagem="", conteudo="", resposta=""):
    return f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <title>Assistente de Documentos</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 760px;
                margin: 60px auto;
                padding: 0 20px;
                color: #1f2937;
            }}
            h1 {{ color: #2563eb; text-align: center; }}
            h2 {{ margin-top: 0; }}
            .caixa, .documentos {{
                margin-top: 28px;
                padding: 28px;
                border-radius: 12px;
                background: #eff6ff;
            }}
            .caixa {{
                border: 2px dashed #93c5fd;
                text-align: center;
            }}
            input, button {{
                margin: 10px;
                padding: 12px;
                font-size: 16px;
            }}
            input[type="text"] {{ width: 70%; }}
            button {{
                border: 0;
                border-radius: 8px;
                background: #2563eb;
                color: white;
                cursor: pointer;
            }}
            .secundario {{ background: #64748b; }}
            .remover {{ background: #dc2626; }}
            .mensagem {{ color: #166534; font-weight: bold; }}
            .conteudo, .resposta {{
                margin-top: 28px;
                padding: 20px;
                border-radius: 12px;
                background: #f8fafc;
                white-space: pre-wrap;
            }}
            .resposta {{ border-left: 5px solid #2563eb; }}
            .documento {{
                display: flex;
                justify-content: space-between;
                align-items: center;
                gap: 16px;
                padding: 16px 0;
                border-top: 1px solid #cbd5e1;
            }}
            .documento:first-of-type {{ border-top: 0; }}
            .documento span {{ color: #475569; font-size: 14px; }}
            .acao {{ display: inline; }}
            .acao button {{ margin: 4px; padding: 8px 12px; font-size: 14px; }}
        </style>
    </head>
    <body>
        <h1>Assistente de Documentos</h1>

        <div class="caixa">
            <h2>Enviar documento</h2>
            <p>Formatos aceitos: texto, Word e PDF.</p>
            <form action="/enviar" method="post" enctype="multipart/form-data">
                <input type="file" name="arquivo"
                       accept=".txt,.docx,.pdf" required>
                <br>
                <button type="submit">Enviar documento</button>
            </form>
            <p class="mensagem">{mensagem}</p>
        </div>

        {documentos_html()}

        <div class="caixa">
            <h2>Faça uma pergunta</h2>
            <form action="/perguntar" method="post">
                <input type="text" name="pergunta"
                       placeholder="Ex.: Faça um resumo do documento"
                       required>
                <br>
                <button type="submit">Perguntar</button>
            </form>
        </div>

        {conteudo}
        {resposta}
    </body>
    </html>
    """


def extrair_texto(caminho):
    extensao = caminho.suffix.lower()

    if extensao == ".txt":
        return caminho.read_text(encoding="utf-8", errors="replace")

    if extensao == ".docx":
        documento = Document(str(caminho))
        partes = [paragrafo.text for paragrafo in documento.paragraphs]

        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.append(" | ".join(celula.text for celula in linha.cells))

        return "\n".join(partes)

    if extensao == ".pdf":
        leitor = PdfReader(str(caminho))
        return "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)

    raise ValueError("Formato de arquivo não aceito.")


def documentos_ativos():
    desativados = carregar_desativados()

    return sorted(
        [
            arquivo
            for arquivo in PASTA_UPLOADS.iterdir()
            if (
                arquivo.is_file()
                and arquivo.suffix.lower() in FORMATOS_ACEITOS
                and arquivo.name not in desativados
            )
        ],
        key=lambda arquivo: arquivo.stat().st_mtime,
    )

def carregar_indice():
    if not ARQUIVO_INDICE.exists():
        return []

    try:
        return json.loads(ARQUIVO_INDICE.read_text(encoding="utf-8"))
    except Exception:
        return []


def salvar_indice(indice):
    ARQUIVO_INDICE.write_text(
        json.dumps(indice, ensure_ascii=False),
        encoding="utf-8",
    )


def dividir_em_trechos(texto):
    texto = " ".join(texto.split())

    if not texto:
        return []

    trechos = []
    passo = TAMANHO_TRECHO - SOBREPOSICAO

    for inicio in range(0, len(texto), passo):
        trecho = texto[inicio : inicio + TAMANHO_TRECHO]

        if trecho:
            trechos.append(trecho)

        if inicio + TAMANHO_TRECHO >= len(texto):
            break

    return trechos


def gerar_embeddings(textos):
    cliente = OpenAI()
    embeddings = []

    for inicio in range(0, len(textos), 100):
        lote = textos[inicio : inicio + 100]
        resultado = cliente.embeddings.create(
            model="text-embedding-3-small",
            input=lote,
        )
        embeddings.extend(item.embedding for item in resultado.data)

    return embeddings


def recriar_indice():
    documentos = documentos_ativos()
    registros = []

    for documento in documentos:
        texto = extrair_texto(documento)

        for trecho in dividir_em_trechos(texto):
            registros.append(
                {
                    "arquivo": documento.name,
                    "trecho": trecho,
                }
            )

    if not registros:
        salvar_indice([])
        return 0, 0

    embeddings = gerar_embeddings(
        [registro["trecho"] for registro in registros]
    )

    for registro, embedding in zip(registros, embeddings):
        registro["embedding"] = embedding

    salvar_indice(registros)
    return len(documentos), len(registros)


def remover_do_indice(nome):
    indice = carregar_indice()
    indice = [
        registro
        for registro in indice
        if registro.get("arquivo") != nome
    ]
    salvar_indice(indice)


def similaridade(vetor_a, vetor_b):
    produto = sum(a * b for a, b in zip(vetor_a, vetor_b))
    norma_a = math.sqrt(sum(a * a for a in vetor_a))
    norma_b = math.sqrt(sum(b * b for b in vetor_b))

    if not norma_a or not norma_b:
        return 0

    return produto / (norma_a * norma_b)


def buscar_trechos_relevantes(pergunta):
    desativados = carregar_desativados()
    indice = carregar_indice()

    registros_ativos = [
        registro
        for registro in indice
        if registro.get("arquivo") not in desativados
        and (PASTA_UPLOADS / registro.get("arquivo", "")).exists()
    ]

    if not registros_ativos:
        return []

    vetor_pergunta = gerar_embeddings([pergunta])[0]

    for registro in registros_ativos:
        registro["pontuacao"] = similaridade(
            vetor_pergunta,
            registro["embedding"],
        )

    return sorted(
        registros_ativos,
        key=lambda registro: registro["pontuacao"],
        reverse=True,
    )[:QUANTIDADE_TRECHOS]

@app.get("/", response_class=HTMLResponse)
def inicio():
    return pagina()


@app.post("/enviar", response_class=HTMLResponse)
async def enviar_documento(arquivo: UploadFile = File(...)):
    nome = Path(arquivo.filename or "arquivo").name
    destino = PASTA_UPLOADS / nome

    if destino.suffix.lower() not in FORMATOS_ACEITOS:
        return pagina("Formato não aceito. Envie texto, Word ou PDF.")

    with destino.open("wb") as arquivo_destino:
        shutil.copyfileobj(arquivo.file, arquivo_destino)

    desativados = carregar_desativados()
    desativados.discard(nome)
    salvar_desativados(desativados)

    try:
        texto = extrair_texto(destino)
    except Exception as erro:
        print(f"Erro ao ler o documento: {erro}")
        return pagina("O arquivo foi enviado, mas não foi possível ler o conteúdo.")

    if not texto.strip():
        return pagina("O arquivo foi enviado, mas não foi encontrado texto nele.")

    conteudo = (
        "<div class='conteudo'><h2>Conteúdo lido</h2>"
        f"{html.escape(texto[:6000])}</div>"
    )

    return pagina(f"Arquivo recebido: {html.escape(nome)}", conteudo)

@app.post("/reindexar", response_class=HTMLResponse)
async def reindexar_documentos():
    if not os.getenv("OPENAI_API_KEY"):
        return pagina(
            "A chave da OpenAI não foi encontrada na configuração."
        )

    try:
        quantidade_documentos, quantidade_trechos = recriar_indice()
    except Exception as erro:
        print(f"Erro ao preparar a biblioteca: {erro}")
        return pagina(
            "Não foi possível preparar a biblioteca agora. "
            "Confira o terminal e tente novamente."
        )

    if not quantidade_trechos:
        return pagina(
            "Não encontrei texto nos documentos ativos para preparar."
        )

    return pagina(
        "Biblioteca preparada com "
        f"{quantidade_documentos} documento(s) e "
        f"{quantidade_trechos} trecho(s)."
    )

@app.post("/alternar", response_class=HTMLResponse)
async def alternar_documento(nome: str = Form(...)):
    nome = Path(nome).name
    arquivo = PASTA_UPLOADS / nome

    if not arquivo.exists():
        return pagina("Documento não encontrado.")

    desativados = carregar_desativados()

    if nome in desativados:
        desativados.remove(nome)
        mensagem = f"Documento reativado: {html.escape(nome)}"
    else:
        desativados.add(nome)
        mensagem = f"Documento desativado: {html.escape(nome)}"

    salvar_desativados(desativados)
    return pagina(mensagem)


@app.post("/remover", response_class=HTMLResponse)
async def remover_documento(nome: str = Form(...)):
    nome = Path(nome).name
    arquivo = PASTA_UPLOADS / nome

    if arquivo.exists():
        arquivo.unlink()

    desativados = carregar_desativados()
    desativados.discard(nome)
    salvar_desativados(desativados)

    return pagina(f"Documento removido: {html.escape(nome)}")

@app.post("/perguntar", response_class=HTMLResponse)
async def perguntar(pergunta: str = Form(...)):
    if not os.getenv("OPENAI_API_KEY"):
        return pagina(
            resposta=(
                "<div class='resposta'>"
                "A chave da OpenAI não foi encontrada na configuração."
                "</div>"
            )
        )

    try:
        trechos = buscar_trechos_relevantes(pergunta)
    except Exception as erro:
        print(f"Erro ao buscar trechos: {erro}")
        return pagina(
            resposta=(
                "<div class='resposta'>"
                "Não foi possível buscar nos documentos agora. "
                "Confira a conexão e tente novamente."
                "</div>"
            )
        )

    if not trechos:
        return pagina(
            resposta=(
                "<div class='resposta'>"
                "A biblioteca ainda não foi preparada. "
                "Clique em “Preparar biblioteca para consulta” "
                "depois de enviar ou alterar documentos."
                "</div>"
            )
        )

    contexto = "\n\n".join(
        (
            f"--- DOCUMENTO: {trecho['arquivo']} ---\n"
            f"{trecho['trecho']}"
        )
        for trecho in trechos
    )

    prompt = f"""
Você é um assistente de consulta de documentos.

Use exclusivamente os trechos recuperados abaixo.
Não use conhecimento externo, suposições ou informações que não estejam
claramente apoiadas pelos documentos.

Responda à pergunta de forma direta e coerente.

Regras de fidelidade:
- Para cada ponto importante, informe entre parênteses o nome do documento
  que apoia a informação.
- Só diga que uma ideia está presente nos dois documentos quando houver
  apoio claro nos dois.
- Se os documentos abordarem o mesmo assunto por perspectivas ou exemplos
  diferentes, chame isso de "tema relacionado", e não de informação idêntica.
- Se uma informação aparecer em apenas um documento, deixe isso explícito.
- Se não houver base suficiente nos trechos, responda:
  "Não encontrei base suficiente nos documentos consultados."
- Só faça comparação entre documentos quando a pergunta pedir comparação.
- Quando a pergunta tratar do pensamento de um autor em vários textos,
  responda com a expressão:
  "Com base nos documentos consultados..."
  Não atribua ao autor uma ideia que não esteja sustentada pelos trechos.

TRECHOS RECUPERADOS:
{contexto}

PERGUNTA:
{pergunta}
"""

    try:
        cliente = OpenAI()
        resultado = cliente.responses.create(
            model="gpt-5.6-luna",
            input=prompt,
        )
        texto_resposta = resultado.output_text
    except Exception as erro:
        print(f"Erro ao consultar a OpenAI: {erro}")
        texto_resposta = (
            "Não foi possível consultar a IA agora. "
            "Confira a conexão e tente novamente."
        )

    fontes = sorted({trecho["arquivo"] for trecho in trechos})
    fontes_html = ", ".join(html.escape(fonte) for fonte in fontes)

    resposta = (
        "<div class='resposta'><h2>Resposta</h2>"
        f"{html.escape(texto_resposta)}"
        "<p><strong>Documentos consultados:</strong> "
        f"{fontes_html}</p></div>"
    )

    return pagina(resposta=resposta)